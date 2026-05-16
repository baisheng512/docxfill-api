"""
DocxFill API 服务 v4 - 支持参数化配置
"""
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, field_validator
import requests as req
import json
import re
import io
import os
import uuid
import time
import threading
from docx import Document
from docx.oxml.ns import qn
from typing import Optional

app = FastAPI(title="DocxFill API", description="Word模板占位符替换服务")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = "/tmp/docxfill_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 文件保留时间（秒），默认1小时
FILE_RETENTION_SECONDS = int(os.environ.get("FILE_RETENTION_SECONDS", 3600))


def cleanup_old_files():
    """后台线程：定期清理超过保留时间的文件"""
    while True:
        try:
            now = time.time()
            for fname in os.listdir(OUTPUT_DIR):
                fpath = os.path.join(OUTPUT_DIR, fname)
                if os.path.isfile(fpath) and (now - os.path.getmtime(fpath)) > FILE_RETENTION_SECONDS:
                    os.remove(fpath)
        except Exception:
            pass
        # 每1小时清理一次
        time.sleep(3600)


# 启动清理线程
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

# 默认配置（当参数未传入时使用）
DEFAULT_PERCENT_FIELDS = {
    "商业用途占比", "住宅用途占比", "城镇住宅用途占比", "商业用地占比", "城镇住宅用地占比",
    "建筑密度", "绿地率", "附加税率", "开发利润率"
}

DEFAULT_THOUSAND_SEP_FIELDS = {
    "总建筑面积", "计容建筑面积", "地下建筑面积", "建筑基底面积",
    "商业建筑面积", "一层商业建筑面积", "二层商业建筑面积", "住宅建筑面积",
    "房屋建筑安装工程费", "勘察设计和前期工程费", "宗地内基础设施建设费",
    "其他费用", "房屋建造成本", "管理费用", "销售费用",
    "投资利息1", "销项税额1", "建安进项", "前期进项",
    "其他费用进项", "销售进项", "抵扣合计", "增值税",
    "增值税附加1", "印花税", "合计税费1", "开发成本1",
    "开发利润1", "总地价"
}

DEFAULT_FIELD_MAPPING = {"其他进项": "其他费用进项", "住宅用途占比": "城镇住宅用途占比"}


class FillRequest(BaseModel):
    template_url: str
    excel_data: str
    percent_fields: Optional[str] = ""
    thousand_sep_fields: Optional[str] = ""
    field_mapping: Optional[str] = ""

    @field_validator('percent_fields', 'thousand_sep_fields', 'field_mapping', mode='before')
    @classmethod
    def none_to_empty(cls, v):
        return "" if v is None else v


def normalize_data(excel_rows):
    """将各种格式的Excel数据转为标准横向字典，支持横向/竖向/键值反转格式"""
    if not excel_rows or not isinstance(excel_rows, list):
        return {}
    
    first_row = excel_rows[0]
    if not isinstance(first_row, dict):
        return {}
    
    # 情况1：标准竖向格式（有"字段名"和"值"列）
    if "字段名" in first_row and "值" in first_row:
        result = {}
        for row in excel_rows:
            key = str(row.get("字段名", "")).strip()
            val = row.get("值")
            if key and val is not None:
                result[key] = val
        return result
    
    # 情况2：readExcel把第二行当表头，导致竖向数据键值反转
    # 特征：每行只有2个键，且有一个键在所有行中都相同
    if len(first_row) == 2 and len(excel_rows) > 2:
        keys = list(first_row.keys())
        # 找出所有行中不变的键
        common_keys = set(keys)
        for row in excel_rows[1:]:
            if isinstance(row, dict):
                common_keys &= set(row.keys())
        
        if len(common_keys) >= 1:
            name_key = list(common_keys)[0]
            val_key = keys[0] if keys[0] != name_key else keys[1]
            
            result = {}
            for row in excel_rows:
                if isinstance(row, dict):
                    field_name = str(row.get(name_key, "")).strip()
                    field_val = row.get(val_key)
                    if field_name and field_val is not None:
                        result[field_name] = field_val
            
            # 验证：字段名合理性检查
            chinese_count = sum(1 for k in result if any('\u4e00' <= c <= '\u9fff' for c in k))
            if chinese_count > len(result) * 0.3:
                return result
    
    # 情况3：标准横向格式
    data = {}
    for row in excel_rows:
        if isinstance(row, dict):
            for k, v in row.items():
                if k and not k.startswith("__EMPTY") and not k.startswith("col_"):
                    data[k] = v
    return data


def parse_config(percent_str, thousand_str, mapping_str):
    """解析传入的配置字符串"""
    percent_str = percent_str or ""
    thousand_str = thousand_str or ""
    mapping_str = mapping_str or ""

    percent_fields = DEFAULT_PERCENT_FIELDS.copy()
    if percent_str and percent_str.strip():
        percent_fields = {f.strip() for f in percent_str.split(",") if f.strip()}

    thousand_sep_fields = DEFAULT_THOUSAND_SEP_FIELDS.copy()
    if thousand_str and thousand_str.strip():
        thousand_sep_fields = {f.strip() for f in thousand_str.split(",") if f.strip()}

    field_mapping = DEFAULT_FIELD_MAPPING.copy()
    if mapping_str and mapping_str.strip():
        for pair in mapping_str.split(","):
            if "=" in pair:
                k, v = pair.split("=", 1)
                field_mapping[k.strip()] = v.strip()

    return percent_fields, thousand_sep_fields, field_mapping


def format_value(key, value, percent_fields, thousand_sep_fields):
    if value is None or (isinstance(value, str) and value.strip() == ""):
        return "资料受限"
    if key in percent_fields:
        try:
            num = float(value)
            return f"{num}%" if num > 1 else f"{round(num * 100, 2)}%"
        except:
            return str(value)
    if key in thousand_sep_fields:
        try:
            num = float(value)
            return f"{int(num):,}" if num == int(num) else f"{num:,.2f}"
        except:
            return str(value)
    if isinstance(value, (int, float)):
        if isinstance(value, float) and value != int(value):
            return f"{value:.2f}"
        return str(int(value)) if isinstance(value, float) else str(value)
    return str(value)


def resolve_key(key, data, field_mapping):
    key_stripped = key.strip()
    if key_stripped in data:
        return key_stripped, data[key_stripped]
    if key_stripped in field_mapping:
        mapped = field_mapping[key_stripped]
        if mapped in data:
            return mapped, data[mapped]
    for data_key in data:
        if data_key.strip() == key_stripped:
            return data_key, data[data_key]
    return None, None


def replace_in_doc(doc, data, percent_fields, thousand_sep_fields, field_mapping):
    count = 0
    pattern = re.compile(r'\{\{(.+?)\}\}')

    for para in doc.paragraphs:
        count += replace_in_para(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += replace_in_para(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping)

    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            try:
                if header and not header.is_linked_to_previous:
                    for para in header.paragraphs:
                        count += replace_in_para(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping)
            except:
                pass
        for footer in [section.footer, section.first_page_footer]:
            try:
                if footer and not footer.is_linked_to_previous:
                    for para in footer.paragraphs:
                        count += replace_in_para(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping)
            except:
                pass
    return count


def highlight_run(run):
    """给run添加黄色高亮"""
    rPr = run._element.get_or_add_rPr()
    highlight = rPr.find(qn('w:highlight'))
    if highlight is None:
        highlight = run._element.makeelement(qn('w:highlight'), {})
        rPr.append(highlight)
    highlight.set(qn('w:val'), 'yellow')


def replace_in_para(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping):
    count = 0
    if not pattern.search(para.text):
        return 0

    for run in para.runs:
        if not run.text:
            continue
        for match in pattern.finditer(run.text):
            key = match.group(1)
            resolved_key, value = resolve_key(key, data, field_mapping)
            if resolved_key is not None:
                formatted = format_value(resolved_key, value, percent_fields, thousand_sep_fields)
                run.text = run.text.replace('{{' + key + '}}', formatted)
                highlight_run(run)
                count += 1

    if pattern.search(para.text):
        count += handle_cross_run(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping)
    return count


def handle_cross_run(para, data, pattern, percent_fields, thousand_sep_fields, field_mapping):
    count = 0
    runs = para.runs
    if not runs:
        return 0

    full_text = ''.join(r.text for r in runs)
    matches = list(pattern.finditer(full_text))
    if not matches:
        return 0

    char_map = []
    for i, r in enumerate(runs):
        for _ in r.text:
            char_map.append(i)

    for match in reversed(matches):
        s, e = match.start(), match.end()
        if s >= len(char_map) or e - 1 >= len(char_map):
            continue
        key = match.group(1)
        resolved_key, value = resolve_key(key, data, field_mapping)
        if resolved_key is None:
            continue
        formatted = format_value(resolved_key, value, percent_fields, thousand_sep_fields)

        fi = char_map[s]
        li = char_map[e - 1]

        if fi == li:
            run = runs[fi]
            rs = sum(len(runs[i].text) for i in range(fi))
            ps = s - rs
            pe = e - rs
            run.text = run.text[:ps] + formatted + run.text[pe:]
            highlight_run(run)
            count += 1
        else:
            fr = runs[fi]
            rs = sum(len(runs[i].text) for i in range(fi))
            ps = s - rs
            fr.text = fr.text[:ps] + formatted
            highlight_run(fr)
            for i in range(fi + 1, li):
                runs[i].text = ''
            lr = runs[li]
            ls = sum(len(runs[i].text) for i in range(li))
            pe = e - ls
            if 0 <= pe <= len(lr.text):
                lr.text = lr.text[pe:]
            else:
                lr.text = ''
            count += 1
    return count


@app.post("/fill")
async def fill_template(request: FillRequest):
    # 解析配置
    percent_fields, thousand_sep_fields, field_mapping = parse_config(
        request.percent_fields, request.thousand_sep_fields, request.field_mapping
    )

    try:
        resp = req.get(request.template_url, timeout=30)
        resp.raise_for_status()
    except Exception as e:
        return {"success": False, "message": f"下载模板失败: {str(e)}", "replaced_count": 0, "file_id": ""}

    try:
        excel_raw = json.loads(request.excel_data)
    except Exception as e:
        return {"success": False, "message": f"JSON解析失败: {str(e)}", "replaced_count": 0, "file_id": ""}

    # 支持两种格式：
    # 格式1：单Sheet数组 [{"字段名":"xx","值":"yy"}, ...]
    # 格式2：多Sheet对象 {"Sheet1名称":[...], "Sheet2名称":[...]}
    if isinstance(excel_raw, dict):
        # 多Sheet格式，合并所有Sheet的数据
        merged_data = {}
        for sheet_name, sheet_rows in excel_raw.items():
            if isinstance(sheet_rows, list) and sheet_rows:
                sheet_data = normalize_data(sheet_rows)
                # 同名字段后出现的Sheet覆盖前面的
                merged_data.update(sheet_data)
        data = merged_data
    elif isinstance(excel_raw, list):
        if not excel_raw:
            return {"success": False, "message": "excel_data必须是非空数组", "replaced_count": 0, "file_id": ""}
        data = normalize_data(excel_raw)
    else:
        return {"success": False, "message": "excel_data格式错误，必须是数组或对象", "replaced_count": 0, "file_id": ""}

    doc = Document(io.BytesIO(resp.content))
    count = replace_in_doc(doc, data, percent_fields, thousand_sep_fields, field_mapping)

    file_id = str(uuid.uuid4())
    path = os.path.join(OUTPUT_DIR, f"{file_id}.docx")
    doc.save(path)

    return {"success": True, "message": f"替换完成，共替换{count}个占位符", "replaced_count": count, "file_id": file_id}


@app.get("/download/{file_id}")
async def download_file(file_id: str):
    if "/" in file_id or ".." in file_id:
        raise HTTPException(status_code=400)
    path = os.path.join(OUTPUT_DIR, f"{file_id}.docx")
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="filled.docx")


@app.post("/upload")
async def upload_docx(request: Request):
    """接收base64编码的docx文件，保存后返回下载URL"""
    try:
        body = await request.json()
        b64_content = body.get("content", "")
        if not b64_content:
            return {"success": False, "message": "缺少content参数", "file_id": "", "download_url": ""}
        import base64
        doc_bytes = base64.b64decode(b64_content)
        fid = str(uuid.uuid4())
        path = os.path.join(OUTPUT_DIR, f"{fid}.docx")
        with open(path, "wb") as f:
            f.write(doc_bytes)
        download_url = f"https://docxfill-api-production.up.railway.app/download/{fid}"
        return {"success": True, "message": "上传成功", "file_id": fid, "download_url": download_url}
    except Exception as e:
        return {"success": False, "message": f"上传失败: {str(e)}", "file_id": "", "download_url": ""}


@app.get("/health")
async def health():
    # 统计当前文件数和总大小
    file_count = 0
    total_size = 0
    for fname in os.listdir(OUTPUT_DIR):
        fpath = os.path.join(OUTPUT_DIR, fname)
        if os.path.isfile(fpath):
            file_count += 1
            total_size += os.path.getsize(fpath)
    return {
        "status": "ok",
        "file_count": file_count,
        "total_size_mb": round(total_size / 1024 / 1024, 2),
        "retention_hours": FILE_RETENTION_SECONDS // 3600
    }


@app.post("/cleanup")
async def manual_cleanup():
    """手动触发清理过期文件"""
    now = time.time()
    removed = 0
    for fname in os.listdir(OUTPUT_DIR):
        fpath = os.path.join(OUTPUT_DIR, fname)
        if os.path.isfile(fpath) and (now - os.path.getmtime(fpath)) > FILE_RETENTION_SECONDS:
            os.remove(fpath)
            removed += 1
    return {"success": True, "message": f"清理完成，删除{removed}个过期文件"}


@app.post("/cleanup-all")
async def cleanup_all():
    """清理所有临时文件（慎用）"""
    removed = 0
    for fname in os.listdir(OUTPUT_DIR):
        fpath = os.path.join(OUTPUT_DIR, fname)
        if os.path.isfile(fpath):
            os.remove(fpath)
            removed += 1
    return {"success": True, "message": f"已清空，删除{removed}个文件"}
